"""Export engine — render a draft to multiple formats.

Supports Markdown, HTML, plain text, CSV, and DOCX. PDF is offered via an
HTML payload the client can print-to-PDF (keeps the backend dependency-light
and works in any environment); DOCX uses python-docx when available.
"""
from __future__ import annotations

import csv
import html
import io
from typing import Any


def to_markdown(draft: dict[str, Any]) -> str:
    lines = [f"# {draft.get('title','Untitled')}", "", draft.get("body", "")]
    if draft.get("headline_options"):
        lines += ["", "## Headline options"] + [f"- {h}" for h in draft["headline_options"]]
    if draft.get("cta_options"):
        lines += ["", "## CTA options"] + [f"- {c}" for c in draft["cta_options"]]
    if draft.get("rationale"):
        lines += ["", "## Rationale", draft["rationale"]]
    if draft.get("posting_guidance"):
        lines += ["", "## Posting guidance", draft["posting_guidance"]]
    return "\n".join(lines)


def to_html(draft: dict[str, Any]) -> str:
    body = html.escape(draft.get("body", "")).replace("\n", "<br>")
    title = html.escape(draft.get("title", "Untitled"))
    extra = ""
    if draft.get("headline_options"):
        items = "".join(f"<li>{html.escape(h)}</li>" for h in draft["headline_options"])
        extra += f"<h2>Headline options</h2><ul>{items}</ul>"
    if draft.get("cta_options"):
        items = "".join(f"<li>{html.escape(c)}</li>" for c in draft["cta_options"])
        extra += f"<h2>CTA options</h2><ul>{items}</ul>"
    return (
        f"<!doctype html><html><head><meta charset='utf-8'><title>{title}</title>"
        "<style>body{font-family:system-ui,Segoe UI,Arial,sans-serif;max-width:720px;"
        "margin:40px auto;line-height:1.6;color:#1f2937}h1{font-size:28px}</style></head>"
        f"<body><h1>{title}</h1><p>{body}</p>{extra}</body></html>"
    )


def to_text(draft: dict[str, Any]) -> str:
    return f"{draft.get('title','Untitled')}\n\n{draft.get('body','')}"


def to_csv(draft: dict[str, Any]) -> str:
    buf = io.StringIO()
    w = csv.writer(buf)
    w.writerow(["field", "value"])
    w.writerow(["title", draft.get("title", "")])
    w.writerow(["body", draft.get("body", "")])
    w.writerow(["headline_options", " | ".join(draft.get("headline_options", []))])
    w.writerow(["cta_options", " | ".join(draft.get("cta_options", []))])
    w.writerow(["rationale", draft.get("rationale", "")])
    w.writerow(["posting_guidance", draft.get("posting_guidance", "")])
    return buf.getvalue()


def to_docx_bytes(draft: dict[str, Any]) -> bytes:
    """Return DOCX bytes. Falls back to a .txt-style payload if python-docx
    is unavailable (caller sets the content-type accordingly)."""
    try:
        from docx import Document
    except Exception:  # pragma: no cover
        return to_text(draft).encode("utf-8")
    doc = Document()
    doc.add_heading(draft.get("title", "Untitled"), level=1)
    for para in draft.get("body", "").split("\n\n"):
        doc.add_paragraph(para)
    if draft.get("headline_options"):
        doc.add_heading("Headline options", level=2)
        for h in draft["headline_options"]:
            doc.add_paragraph(h, style="List Bullet")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


EXPORTERS = {
    "markdown": ("text/markdown", to_markdown, "md"),
    "html": ("text/html", to_html, "html"),
    "text": ("text/plain", to_text, "txt"),
    "csv": ("text/csv", to_csv, "csv"),
}
